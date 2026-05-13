"""Generate the lean-SMB software licence proposal as a .docx file."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

OUT_PATH = Path(__file__).resolve().parent.parent / "docs" / "proposals" / "lean-smb-proposal.docx"
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xB8, 0x86, 0x1B)


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text: str, size: int = 13, color: RGBColor = INK, space_before: int = 8) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_paragraph(doc, text: str, size: int = 10, bold: bool = False, italic: bool = False, color: RGBColor = INK, align=None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_table(doc, rows, header: bool = True, col_widths=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            if header and r_idx == 0:
                shade_cell(cell, "F2EFE6")
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if header and r_idx == 0:
                run.bold = True
                run.font.color.rgb = INK
            if col_widths and c_idx < len(col_widths):
                cell.width = col_widths[c_idx]
    return table


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# Title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
tr = title.add_run("SOFTWARE LICENCE PROPOSAL")
tr.bold = True
tr.font.size = Pt(18)
tr.font.color.rgb = INK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(10)
sr = sub.add_run("Production & Repair Order Management Platform")
sr.font.size = Pt(11)
sr.font.color.rgb = MUTED
sr.italic = True

meta_table = doc.add_table(rows=1, cols=2)
meta_table.autofit = True
meta_left = meta_table.rows[0].cells[0]
meta_right = meta_table.rows[0].cells[1]
meta_left.text = ""
meta_right.text = ""

for label, value in [("Prepared for:", "[Client Name]"), ("Prepared by:", "North East Engineering — Software Division")]:
    p = meta_left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    rb = p.add_run(f"{label} ")
    rb.bold = True
    rb.font.size = Pt(10)
    rv = p.add_run(value)
    rv.font.size = Pt(10)

for label, value in [("Date:", "14 May 2026"), ("Validity:", "30 days from issue")]:
    p = meta_right.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rb = p.add_run(f"{label} ")
    rb.bold = True
    rb.font.size = Pt(10)
    rv = p.add_run(value)
    rv.font.size = Pt(10)

# Strip default first empty paragraph on each cell
for cell in (meta_left, meta_right):
    if cell.paragraphs and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
        p._element.getparent().remove(p._element)

# 1. Executive Summary
add_heading(doc, "1. Executive Summary", space_before=10)
add_paragraph(
    doc,
    "We propose a one-time, perpetual licence to deploy our purpose-built Production & Repair Order "
    "Management Platform at your single-site workshop. The platform replaces approximately 80% of the "
    "day-to-day functionality of Epicor Kinetic Production at a fraction of the cost, with zero "
    "implementation overhead.",
)

# 2. What You Receive
add_heading(doc, "2. What You Receive")
modules = [
    ("Module", "Description"),
    ("Repair Order Management", "Full RO lifecycle from intake to handover, 10-stage Kanban workflow"),
    ("Job Task Scheduling", "Station-based assignment with gate controls (Draft / Approval / Chassis)"),
    ("Templates & Versioning", "Reusable RO definitions with estimated hours per operation"),
    ("Time Tracking & Variance", "Per-task clock in/out with estimate-vs-actual reporting"),
    ("Customer Approval Workflow", "Digital quote sign-off captured before work commences"),
    ("Chassis Inventory", "Serialised chassis tracking and allocation"),
    ("Technician Mobile UI", "Shop-floor optimised interface for task execution"),
    ("Role-Based Access (8 roles)", "Supervisor, Sales, Drafter, QC, Technician, and more"),
    ("Two Themes (Light + SaaS)", "Branded UI suitable for office and shop-floor displays"),
    ("Reporting Dashboard", "Built-in operational reports"),
]
add_table(doc, modules, col_widths=[Cm(5.0), Cm(11.5)])

# 3. Investment
add_heading(doc, "3. Investment")
investment = [
    ("Item", "Amount (AUD)", "Amount (INR)"),
    ("One-time perpetual licence fee", "$45,000", "₹25,00,000"),
    ("Hosting & infrastructure", "Excluded — by client", "Excluded"),
    ("Annual maintenance & support", "Optional, quoted separately", "Optional"),
]
add_table(doc, investment, col_widths=[Cm(8.5), Cm(4.5), Cm(3.5)])
add_paragraph(
    doc,
    "Payment terms: 50% on contract signing · 50% on go-live (within 30 days of signing).",
    italic=True,
    color=MUTED,
)

# 4. Licence Scope
add_heading(doc, "4. Licence Scope")
for line in [
    "Sites: 1 (one) physical workshop location",
    "Named users: Up to 15",
    "Term: Perpetual — no annual renewal required",
    "Upgrades: Bug fixes for 12 months included; feature releases optional",
    "Source code: Not included (executable licence only)",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line)
    r.font.size = Pt(10)

# 5. Exclusions
add_heading(doc, "5. What is Not Included")
for line in [
    "Cloud or on-premise hosting infrastructure",
    "Custom feature development, integrations, or report changes",
    "Data migration from existing systems",
    "On-site training (remote handover session included; on-site billed at AUD 1,500/day)",
    "Ongoing maintenance, monitoring, or 24×7 support",
    "MRP, General Ledger, AP/AR, Purchasing, or Forecasting modules",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line)
    r.font.size = Pt(10)
add_paragraph(
    doc,
    "Any item above is available as a separately quoted engagement at AUD 150/hour.",
    italic=True,
    color=MUTED,
)

# 6. Comparison
add_heading(doc, "6. Why This Beats Epicor Kinetic")
comparison = [
    ("", "This Proposal", "Epicor Kinetic Production"),
    ("One-time cost", "AUD 45,000", "AUD 250,000+"),
    ("Year 1 total (incl. implementation)", "AUD 45,000", "AUD 350,000 – 500,000"),
    ("Annual recurring", "Optional", "AUD 80,000 – 150,000"),
    ("Time to go-live", "2 – 4 weeks", "6 – 12 months"),
    ("Built for heavy-vehicle repair", "Yes — out of the box", "Requires customisation"),
]
add_table(doc, comparison, col_widths=[Cm(6.5), Cm(5.0), Cm(5.0)])

# 7. Acceptance
add_heading(doc, "7. Acceptance")
add_paragraph(doc, "Signed for and on behalf of the Client:")
doc.add_paragraph()  # spacing

sig_table = doc.add_table(rows=2, cols=2)
sig_table.autofit = True
sig_rows = [
    ("Name: ____________________________", "Title: ____________________________"),
    ("Signature: _______________________", "Date: ____________________________"),
]
for r_idx, row in enumerate(sig_rows):
    for c_idx, val in enumerate(row):
        cell = sig_table.rows[r_idx].cells[c_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(val)
        run.font.size = Pt(10)

# Footer
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(14)
fr = footer_p.add_run("This proposal is confidential and intended solely for the named recipient.")
fr.italic = True
fr.font.size = Pt(8)
fr.font.color.rgb = MUTED

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
